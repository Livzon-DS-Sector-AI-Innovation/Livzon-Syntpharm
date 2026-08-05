"""Generate S6 Template Fixture with aligned anchors."""
from docx import Document
from docx.shared import Inches
from PIL import Image
import os

def generate():
    doc = Document()
    
    # 1. S6 Numbered Heading
    doc.add_heading('3.2.S.6 包装系统', level=1)
    doc.add_paragraph('本章节描述药品的包装系统细节。')
    
    # 2. Anchor Paragraph for "包装形式" (Mapping: paragraph, keyword='包装形式')
    doc.add_paragraph('包装形式：')
    
    # 3. Table for other mappings
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = '包材类型'
    table.cell(0, 1).text = ''
    table.cell(1, 0).text = '厂内名称'
    table.cell(1, 1).text = ''
    table.cell(2, 0).text = '包材生产商'
    table.cell(2, 1).text = ''

    # 4. Real Image (Media >= 1)
    img = Image.new('RGB', (100, 100), color='blue')
    img_path = '/tmp/fixture_img.png'
    img.save(img_path)
    doc.add_picture(img_path, width=Inches(1.0))
    
    # 5. Appendix Anchors (Matching _insert_image_at_appendix logic)
    # Logic: if appendix_slot not in text or "\t" in para.text: continue
    #        if text == appendix_slot: continue
    # So we need: "附录1 营业执照位置" (Contains "附录1", not equal to "附录1", no tab)
    doc.add_heading('附录1 营业执照位置', level=2)
    doc.add_paragraph('此处插入营业执照图片。')
    
    doc.add_heading('附录2 CDE公示位置', level=2)
    
    os.makedirs(os.path.dirname(__file__) + '/dossier_splits', exist_ok=True)
    output_path = os.path.dirname(__file__) + '/dossier_splits/s6_template.docx'
    doc.save(output_path)
    print(f'Fixture generated at: {output_path}')
    
    # Verify Media
    from zipfile import ZipFile
    with ZipFile(output_path) as z:
        media = [f for f in z.namelist() if 'media' in f]
        print(f'Template Media Count: {len(media)}')

if __name__ == "__main__":
    generate()
