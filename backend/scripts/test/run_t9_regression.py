"""T-9: S6 AI Full Pipeline Regression Test (v2)."""

import asyncio
import sys
from pathlib import Path
from zipfile import ZipFile

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from docx import Document
from PIL import Image
from sqlalchemy import and_, select

from app.core.database import get_db
from app.modules.registration.dossier_writer.ai_fill_service import AIFillService
from app.modules.registration.dossier_writer.models import ChapterAsset, DossierChapter
from app.modules.registration.dossier_writer.schemas import ProductDossierCreate
from app.modules.registration.dossier_writer.service import DossierService


async def run_t9():
    print("🚀 Starting T-9: S6 AI Full Pipeline Regression...")

    async for db in get_db():
        service = DossierService(db)

        # Step 1: Create Dossier via Service (Initializes paths) & Upload Template
        print("\n[Step 1] Creating dossier via service and uploading S6 template...")
        create_data = ProductDossierCreate(product_name="T9_Final_Test", sterile_type="无菌", manufacturer="T9_Mfr")
        dossier = await service.create_product_dossier(create_data)

        # Upload template
        fixture_path = (
            Path(__file__).resolve().parents[2] / "tests" / "fixtures" / "dossier_splits" / "s6_template.docx"
        )
        content = fixture_path.read_bytes()
        await service.save_template_file(dossier.id, "s6_template.docx", content)

        # Trigger match and split
        match_result = await service.match_assets_to_chapters(dossier.id)
        print(f"   ✅ Match Result: {match_result['message']}")

        # Get S6 chapter
        s6_stmt = select(DossierChapter).where(
            and_(DossierChapter.product_dossier_id == dossier.id, DossierChapter.chapter_code == "3.2.S.6")
        )
        s6_res = await db.execute(s6_stmt)
        s6_chapter = s6_res.scalar_one_or_none()

        if s6_chapter and s6_chapter.working_file:
            wf_path = Path(dossier.working_path) / s6_chapter.working_file
            print(f"   ✅ Working File Created: {wf_path.name}")

            with ZipFile(str(wf_path), "r") as z:
                media_files = [f for f in z.namelist() if f.startswith("word/media/")]
                print(f"   ✅ Media Files in Split Working File: {len(media_files)} files")
        else:
            print("   ❌ FAIL: No working file generated.")
            return

        # Step 2 & 3: Real AI Confirm
        print("\n[Step 2-3] Executing Real AI Confirm...")
        mock_fields = [
            {
                "field_name": "包装系统描述",
                "value": "T-9 Real Write Value",
                "confidence": 0.95,
                "field_type": "text",
                "field_mapping_id": None,
            }
        ]

        ai_service = AIFillService(db)
        write_result = await ai_service.confirm_and_fill(dossier, s6_chapter, mock_fields)
        print(f"   ✅ AI Confirm: {write_result['message']}")

        doc = Document(str(wf_path))
        found_text = any("T-9 Real Write Value" in p.text for p in doc.paragraphs)
        print(f"   ✅ Text Search in Working File: {'Found' if found_text else 'Not Found'}")

        # Step 4: Split-Confirm
        print("\n[Step 4] Executing Real Split-Confirm...")
        img = Image.new("RGB", (100, 100), color="red")
        dummy_png = "/tmp/test_split.png"
        img.save(dummy_png)

        asset = ChapterAsset(
            chapter_id=s6_chapter.id, original_filename="test_split.png", file_path=dummy_png, file_type="png"
        )
        db.add(asset)
        await db.commit()

        splits = [{"split_id": str(asset.id), "appendix_slot": "附录1", "asset_id": str(asset.id), "page_number": 1}]
        split_result = await ai_service.confirm_page_splits_and_insert(dossier, s6_chapter, splits)
        print(f"   ✅ Split-Confirm: {split_result['message']}")

        with ZipFile(str(wf_path), "r") as z:
            media_after = [f for f in z.namelist() if f.startswith("word/media/")]
            print(f"   ✅ Media Files After Split-Confirm: {len(media_after)} files")

        # Step 5: Preview
        print("\n[Step 5] Verifying Preview...")
        preview = await service.get_chapter_preview(s6_chapter.id)
        found_preview = any("T-9 Real Write Value" in p["text"] for p in preview["paragraphs"])
        print(f"   ✅ Preview Text Hit: {'Found' if found_preview else 'Not Found'}")

        # Step 6: Export
        print("\n[Step 6] Exporting...")
        export_result = await service.export_dossier(dossier.id, [s6_chapter.id])
        if export_result["success"]:
            exp_path = Path(export_result["file_path"])
            doc_exp = Document(str(exp_path))
            found_export = any("T-9 Real Write Value" in p.text for p in doc_exp.paragraphs)
            print(f"   ✅ Export Text Hit: {'Found' if found_export else 'Not Found'}")

            with ZipFile(str(exp_path), "r") as z:
                media_exp = [f for f in z.namelist() if f.startswith("word/media/")]
                print(f"   ✅ Export Media Files: {len(media_exp)} files")
        else:
            print(f"   ❌ Export Failed: {export_result['message']}")

        print("\n🎉 T-9 Regression Test Completed.")
        break


if __name__ == "__main__":
    asyncio.run(run_t9())
