"""素材文本提取器 - 将各种格式的素材统一转为纯文本，供 AI 解析"""

import logging
from pathlib import Path
from typing import Any

from docx import Document

_logger = logging.getLogger(__name__)


class AssetTextExtractor:
    """素材文本提取器"""

    @staticmethod
    def extract(file_path: Path) -> dict[str, Any]:
        """统一提取接口：根据文件类型选择提取方式

        Returns:
            {
                "text": "提取的全文",
                "paragraphs": [{"index": 0, "text": "..."}],
                "tables": [{"index": 0, "rows": N, "cols": M, "data": [[...]]}],
                "page_count": int (仅 PDF),
                "page_texts": [{"page": 1, "text": "..."}] (仅 PDF),
                "error": str (如有错误)
            }
        """
        # 检查文件是否存在
        if not file_path.exists():
            return {"text": "", "error": f"素材文件不存在，请重新上传: {file_path.name}"}

        suffix = file_path.suffix.lower()

        if suffix == ".docx":
            return AssetTextExtractor._extract_docx(file_path)
        elif suffix == ".doc":
            return AssetTextExtractor._extract_doc(file_path)
        elif suffix == ".pdf":
            return AssetTextExtractor._extract_pdf(file_path)
        elif suffix == ".xlsx":
            return AssetTextExtractor._extract_xlsx(file_path)
        elif suffix == ".xls":
            return AssetTextExtractor._extract_xls(file_path)
        elif suffix in (".txt", ".csv"):
            return AssetTextExtractor._extract_text(file_path)
        else:
            return {"text": "", "error": f"不支持的文件类型: {suffix}"}

    @staticmethod
    def _extract_docx(file_path: Path) -> dict[str, Any]:
        """从 docx 提取段落和表格"""
        try:
            doc = Document(str(file_path))

            paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    paragraphs.append({"index": i, "text": text})

            tables = []
            for t_idx, table in enumerate(doc.tables):
                rows_data = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows_data.append(cells)
                tables.append(
                    {
                        "index": t_idx,
                        "rows": len(table.rows),
                        "cols": len(table.columns),
                        "data": rows_data,
                    }
                )

            full_text = "\n".join(str(p["text"]) for p in paragraphs)
            return {"text": full_text, "paragraphs": paragraphs, "tables": tables}
        except Exception as e:
            return {"text": "", "error": f"docx 提取失败: {str(e)}"}

    @staticmethod
    def _ensure_ocr_service(timeout: int = 180) -> Any:
        """确保 OCR 服务已初始化，如未初始化则尝试自动初始化一次

        Args:
            timeout: OCR 初始化超时时间（秒），默认 180 秒

        Returns:
            OCRService 实例，或 None（如果初始化失败或超时）
        """
        import threading

        from app.shared import ocr_service as ocr_module

        # 如果已初始化，直接返回
        if ocr_module._ocr_service is not None:
            return ocr_module._ocr_service

        # 如果正在初始化，返回错误
        if ocr_module._ocr_initializing:
            _logger.warning("OCR 服务正在初始化中，请稍后重试")
            return None

        # 尝试自动初始化，带超时控制
        _logger.info(f"OCR 服务未初始化，尝试自动初始化（超时: {timeout}秒）...")

        init_result = [None]
        init_exception = [None]

        def init_worker() -> None:
            try:
                ocr_module.init_ocr()
                init_result[0] = ocr_module._ocr_service  # type: ignore[assignment]
            except Exception as e:
                init_exception[0] = e  # type: ignore[call-overload]

        # 在独立线程中初始化
        init_thread = threading.Thread(target=init_worker)
        init_thread.start()
        init_thread.join(timeout=timeout)

        if init_thread.is_alive():
            _logger.error(f"OCR 服务初始化超时（{timeout}秒）")
            return None

        if init_exception[0] is not None:
            _logger.error(f"OCR 服务初始化失败: {init_exception[0]}")
            return None

        if init_result[0] is not None:
            _logger.info("OCR 服务自动初始化成功")
            return init_result[0]

        _logger.error("OCR 服务初始化后仍不可用")
        return None

    @staticmethod
    def _extract_doc(file_path: Path) -> dict[str, Any]:
        """从 .doc 提取：先转 docx 再提取（使用共享文件转换服务）"""
        from app.shared.file_conversion import get_file_conversion

        service = get_file_conversion()
        docx_path = service.convert_to_docx(file_path)
        if docx_path is None or not docx_path.exists():
            return {"text": "", "error": "LibreOffice 未安装或转换失败，请上传 .docx 格式"}
        return AssetTextExtractor._extract_docx(docx_path)

    @staticmethod
    def _extract_pdf(file_path: Path) -> dict[str, Any]:
        """从 PDF 提取：优先用 pdfplumber，失败则用 OCR"""
        # 尝试 pdfplumber（对文字型 PDF 更快更准）
        try:
            import pdfplumber

            with pdfplumber.open(str(file_path)) as pdf:
                page_texts = []
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    page_texts.append({"page": i + 1, "text": text.strip()})

                full_text = "\n\n".join(str(p["text"]) for p in page_texts)
                if full_text.strip():
                    return {
                        "text": full_text,
                        "page_count": len(pdf.pages),
                        "page_texts": page_texts,
                    }
        except Exception:
            _logger.warning("PDF extraction with pdfplumber failed, falling back to OCR")

        # 回退到 OCR（对扫描件）
        _logger.info("PDF 无文本内容，尝试使用 OCR...")
        ocr_service = AssetTextExtractor._ensure_ocr_service(timeout=180)
        if ocr_service is None:
            return {
                "text": "",
                "error": "该 PDF 为扫描件，需要 OCR，但当前 OCR 服务不可用或初始化超时（180秒）。请上传可复制文本的 PDF、Word 或 Excel 文件。",
            }

        try:
            # 使用线程包裹 OCR 提取，防止无限阻塞
            import threading

            ocr_timeout = 180  # 单个文件 OCR 提取最大耗时（秒）

            ocr_result = {"markdown": None, "structure": None, "error": None}

            def ocr_worker() -> None:
                try:
                    ocr_result["markdown"] = ocr_service.extract_markdown(file_path)
                    ocr_result["structure"] = ocr_service.extract_structure(file_path)
                except Exception as e:
                    ocr_result["error"] = str(e)  # type: ignore[assignment]

            worker_thread = threading.Thread(target=ocr_worker, daemon=True)
            worker_thread.start()
            worker_thread.join(timeout=ocr_timeout)

            if worker_thread.is_alive():
                _logger.error(f"OCR 提取超时（{ocr_timeout}秒）: {file_path}")
                return {
                    "text": "",
                    "error": f"该 PDF 为扫描件，OCR 提取超时（{ocr_timeout}秒）。请上传可复制文本的 PDF、Word 或 Excel 文件。",
                }

            if ocr_result["error"]:
                _logger.error(f"OCR 提取失败: {ocr_result['error']}")
                return {"text": "", "error": f"PDF OCR 提取失败: {ocr_result['error']}"}

            markdown_text = ocr_result["markdown"]
            structure = ocr_result["structure"]

            return {
                "text": markdown_text,
                "page_count": 1,
                "page_texts": [{"page": 1, "text": markdown_text}],
                "structure": structure,
            }
        except Exception as e:
            _logger.error(f"OCR 提取失败: {e}")
            return {"text": "", "error": f"PDF OCR 提取失败: {str(e)}"}

    @staticmethod
    def _extract_text(file_path: Path) -> dict[str, Any]:
        """从纯文本文件提取"""
        try:
            text = file_path.read_text(encoding="utf-8", errors="replace")
            paragraphs = [{"index": i, "text": line} for i, line in enumerate(text.splitlines()) if line.strip()]
            return {"text": text, "paragraphs": paragraphs, "tables": []}
        except Exception as e:
            return {"text": "", "error": f"文本提取失败: {str(e)}"}

    @staticmethod
    def _extract_xlsx(file_path: Path) -> dict[str, Any]:
        """从 xlsx 提取表格数据"""
        try:
            import openpyxl

            wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
            tables = []
            all_text = []

            for sheet_idx, sheet_name in enumerate(wb.sheetnames):
                ws = wb[sheet_name]
                rows_data = []
                for row in ws.iter_rows(values_only=True):
                    cells = [str(cell) if cell is not None else "" for cell in row]
                    if any(cells):
                        rows_data.append(cells)
                        all_text.append("\t".join(cells))

                if rows_data:
                    tables.append(
                        {
                            "index": sheet_idx,
                            "sheet_name": sheet_name,
                            "rows": len(rows_data),
                            "cols": max(len(r) for r in rows_data) if rows_data else 0,
                            "data": rows_data,
                        }
                    )

            wb.close()
            full_text = "\n".join(all_text)
            return {"text": full_text, "paragraphs": [], "tables": tables}
        except Exception as e:
            return {"text": "", "error": f"xlsx 提取失败: {str(e)}"}

    @staticmethod
    def _extract_xls(file_path: Path) -> dict[str, Any]:
        """从 xls 提取表格数据"""
        try:
            import pandas as pd

            excel_file = pd.ExcelFile(str(file_path), engine="xlrd")
            tables = []
            all_text = []

            for sheet_idx, sheet_name in enumerate(excel_file.sheet_names):
                df = excel_file.parse(sheet_name)
                rows_data = []
                for _, row in df.iterrows():
                    cells = [str(cell) if pd.notna(cell) else "" for cell in row]
                    if any(cells):
                        rows_data.append(cells)
                        all_text.append("\t".join(cells))

                if rows_data:
                    tables.append(
                        {
                            "index": sheet_idx,
                            "sheet_name": sheet_name,
                            "rows": len(rows_data),
                            "cols": max(len(df.columns), len(rows_data[0])) if rows_data else 0,
                            "data": rows_data,
                        }
                    )

            full_text = "\n".join(all_text)
            return {"text": full_text, "paragraphs": [], "tables": tables}
        except Exception as e:
            return {"text": "", "error": f"xls 提取失败: {str(e)}"}

    @staticmethod
    def pdf_page_to_image(file_path: Path, page_number: int, dpi: int = 200) -> Path | None:
        """将 PDF 指定页转为图片，返回图片路径"""
        try:
            from pdf2image import convert_from_path

            images = convert_from_path(
                str(file_path),
                dpi=dpi,
                first_page=page_number,
                last_page=page_number,
            )
            if not images:
                return None

            img_path = file_path.parent / f"{file_path.stem}_page{page_number}.png"
            images[0].save(str(img_path), "PNG")
            return img_path
        except Exception:
            return None
